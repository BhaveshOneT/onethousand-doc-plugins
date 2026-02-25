#!/usr/bin/env python3
"""
Generate branded hackathon debrief DOCX documents from JSON content.

Translates the TypeScript DocxGenerator class from the hackathon-web app into
Python using python-docx.  Produces a two-section document:

  Section 1 -- Full-page green title page built from a single-cell table.
  Section 2 -- CONTINUOUS break, 1-inch margins, TOC + content sections with
               markdown conversion, headers, footers, and inline images.

Usage:
    python generate_debrief_doc.py \\
        --content /tmp/debrief_content.json \\
        --logo-dir <path_to_logos_dir> \\
        --output /path/to/output.docx
"""

from __future__ import annotations

import argparse
import base64
import io
import json
import os
import re
import struct
import sys
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple, Union

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, Twips, RGBColor

# Optional PIL for inline base64 image sizing
try:
    from PIL import Image as PILImage
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

# ---------------------------------------------------------------------------
# Brand constants (matches TypeScript exactly)
# ---------------------------------------------------------------------------

BRAND_COLORS = {
    'sharpGreen': '19A960',
    'ash': '2F2F2F',
    'neuralLime': 'D5F89E',
    'ceruleanBlue': '829DB6',
    'deepBlack': '000000',
    'pureWhite': 'FFFFFF',
}

FONT_FAMILIES = {
    'heading': 'Amsi Pro Narw Black',
    'body': 'Akkurat LL',
    'display': 'Amsi Pro Narw Black',
    'display_title': 'Wavetable',
    'mono': 'Consolas',
}

PAGE_SIZE_A4 = {
    'width': Mm(210),   # 11906 twips
    'height': Mm(297),  # 16838 twips
}

TABLE_ROW_HEIGHT_TWIPS = 17600

SECTION_ORDER: List[str] = [
    'background',
    'hackathon_structure',
    'challenge',
    'goal',
    'data',
    'approach',
    'results',
    'canvas',
    'user_flow',
    'conclusion',
]

SECTION_TITLES: Dict[str, Dict[str, str]] = {
    'participants':        {'de': 'Teilnehmer',            'en': 'Participants'},
    'background':          {'de': 'Hintergrund',           'en': 'Background'},
    'hackathon_structure': {'de': 'Hackathon',             'en': 'Hackathon'},
    'challenge':           {'de': 'Herausforderung',       'en': 'Challenge'},
    'goal':                {'de': 'Ziel',                  'en': 'Goal'},
    'data':                {'de': 'Daten',                 'en': 'Data'},
    'approach':            {'de': 'Ansatz',                'en': 'Approach'},
    'results':             {'de': 'Ergebnisse',            'en': 'Results'},
    'canvas':              {'de': 'AI Breakthrough Canvas','en': 'AI Breakthrough Canvas'},
    'user_flow':           {'de': 'Benutzerfluss',         'en': 'User Flow'},
    'conclusion':          {'de': 'Fazit',                 'en': 'Conclusion'},
}


# ---------------------------------------------------------------------------
# Title page adaptive profiles
# ---------------------------------------------------------------------------

@dataclass
class TitlePageProfile:
    """Adaptive sizing parameters for the cover page."""
    logoSize: int
    logoAfter: int
    titleSpacerBefore: int
    titleSize: int          # half-points in TS; converted to Pt when used
    debriefAfter: int
    subtitleSize: int
    subtitleAfter: int
    companyHeaderSize: int
    companyAfter: int
    participantSize: int
    participantAfter: int
    oneThousandBefore: int
    oneThousandAfter: int
    footerBefore: int
    footerTitleSize: int
    dateSize: int


PROFILE_DEFAULT = TitlePageProfile(
    logoSize=54, logoAfter=300, titleSpacerBefore=500, titleSize=72,
    debriefAfter=240, subtitleSize=44, subtitleAfter=320,
    companyHeaderSize=20, companyAfter=70, participantSize=20, participantAfter=20,
    oneThousandBefore=20, oneThousandAfter=20,
    footerBefore=420, footerTitleSize=24, dateSize=20,
)

PROFILE_MEDIUM = TitlePageProfile(
    logoSize=48, logoAfter=240, titleSpacerBefore=300, titleSize=64,
    debriefAfter=180, subtitleSize=40, subtitleAfter=240,
    companyHeaderSize=18, companyAfter=50, participantSize=18, participantAfter=14,
    oneThousandBefore=14, oneThousandAfter=14,
    footerBefore=300, footerTitleSize=22, dateSize=18,
)

PROFILE_COMPACT = TitlePageProfile(
    logoSize=42, logoAfter=180, titleSpacerBefore=200, titleSize=56,
    debriefAfter=120, subtitleSize=36, subtitleAfter=180,
    companyHeaderSize=16, companyAfter=40, participantSize=16, participantAfter=10,
    oneThousandBefore=10, oneThousandAfter=10,
    footerBefore=200, footerTitleSize=20, dateSize=16,
)


def _get_title_page_profile(
    participant_count: int,
    company_name_length: int,
) -> TitlePageProfile:
    """Select adaptive sizing profile based on participant count and company name length."""
    if participant_count > 20 or company_name_length > 40:
        return PROFILE_COMPACT
    if participant_count > 14 or company_name_length > 30:
        return PROFILE_MEDIUM
    return PROFILE_DEFAULT


# ---------------------------------------------------------------------------
# TOC entry builder (mirrors toc-structure.ts)
# ---------------------------------------------------------------------------

@dataclass
class TocEntry:
    key: str
    title: str
    level: int          # 1 or 2
    section_id: Optional[str] = None
    number_label: str = ''


def _build_toc_entries(
    sections: List[Dict[str, Any]],
    language: str,
    use_cases: Optional[List[Dict[str, Any]]] = None,
) -> List[TocEntry]:
    """Build flat TOC entries — all level-1, sequential numbering."""
    by_id: Dict[str, Dict[str, Any]] = {s['id']: s for s in sections}

    def _title(sid: str) -> str:
        sec = by_id.get(sid)
        if sec and sec.get('title', '').strip():
            return sec['title'].strip()
        return SECTION_TITLES.get(sid, {}).get(language, sid)

    entries: List[TocEntry] = []

    # Build entries from sections that actually exist in the content JSON,
    # following the canonical SECTION_ORDER.
    for sid in SECTION_ORDER:
        if sid not in by_id:
            continue
        entries.append(TocEntry(key=sid, title=_title(sid), level=1, section_id=sid))

    # Apply sequential numbering — all level-1
    for i, entry in enumerate(entries, 1):
        entry.number_label = str(i)

    return entries


# ---------------------------------------------------------------------------
# Markdown block parser (simplified, matching TS parseMarkdownBlocks)
# ---------------------------------------------------------------------------

@dataclass
class MdHeading:
    type: str = 'heading'
    level: int = 2
    text: str = ''

@dataclass
class MdParagraph:
    type: str = 'paragraph'
    text: str = ''

@dataclass
class MdListItem:
    text: str = ''
    indent: int = 0  # 0 = top-level, 1 = sub-item, 2 = sub-sub-item

@dataclass
class MdList:
    type: str = 'list'
    ordered: bool = False
    items: List[Any] = field(default_factory=list)  # List[str] or List[MdListItem]

@dataclass
class MdImage:
    type: str = 'image'
    alt: str = ''
    url: str = ''

@dataclass
class MdTable:
    type: str = 'table'
    headers: List[str] = field(default_factory=list)
    rows: List[List[str]] = field(default_factory=list)

MdBlock = Union[MdHeading, MdParagraph, MdList, MdImage, MdTable]


def _strip_code_blocks(text: str) -> str:
    """Remove fenced code blocks (```...```)."""
    return re.sub(r'```[\s\S]*?```', '', text).strip()


def _strip_redundant_heading(title: str, content: str) -> str:
    """Remove a leading ## heading if it duplicates the section title."""
    lines = content.split('\n')
    i = 0
    while i < len(lines) and lines[i].strip() == '':
        i += 1
    if i < len(lines):
        m = re.match(r'^##\s+(.+)$', lines[i])
        if m and m.group(1).strip().lower() == title.strip().lower():
            lines.pop(i)
            while i < len(lines) and lines[i].strip() == '':
                lines.pop(i)
    return '\n'.join(lines).strip()


def _parse_markdown_blocks(markdown: str) -> List[MdBlock]:
    """Parse markdown text into structured blocks."""
    if not markdown:
        return []

    # Strip fenced code blocks
    text = _strip_code_blocks(markdown)
    lines = text.split('\n')
    blocks: List[MdBlock] = []
    i = 0

    def _is_table_candidate(idx: int) -> bool:
        if idx + 1 >= len(lines):
            return False
        header = lines[idx].strip()
        divider = lines[idx + 1].strip()
        if '|' not in header:
            return False
        return bool(re.match(r'^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$', divider))

    def _split_table_row(line: str) -> List[str]:
        row = line.strip()
        if row.startswith('|'):
            row = row[1:]
        if row.endswith('|'):
            row = row[:-1]
        return [c.strip() for c in row.split('|')]

    while i < len(lines):
        line = lines[i]
        trimmed = line.strip()

        if not trimmed:
            i += 1
            continue

        # Heading
        hm = re.match(r'^(#{1,6})\s+(.+)$', trimmed)
        if hm:
            blocks.append(MdHeading(level=len(hm.group(1)), text=hm.group(2).strip()))
            i += 1
            continue

        # Image  ![alt](url)
        if trimmed.startswith('!['):
            alt_end = trimmed.find('](')
            if alt_end >= 2 and trimmed.endswith(')'):
                alt = trimmed[2:alt_end].strip()
                url = trimmed[alt_end + 2:-1].strip()
                if url:
                    blocks.append(MdImage(alt=alt, url=url))
                    i += 1
                    continue

        # Table
        if _is_table_candidate(i):
            headers = _split_table_row(lines[i])
            i += 2  # skip header + separator
            rows: List[List[str]] = []
            while i < len(lines) and '|' in lines[i].strip() and lines[i].strip():
                rows.append(_split_table_row(lines[i]))
                i += 1
            blocks.append(MdTable(headers=headers, rows=rows))
            continue

        # Ordered list
        if re.match(r'^\d+\.\s+', trimmed):
            items: List[str] = []
            while i < len(lines):
                lt = lines[i].strip()
                if re.match(r'^\d+\.\s+', lt):
                    items.append(re.sub(r'^\d+\.\s+', '', lt))
                    i += 1
                elif lt == '':
                    peek = i + 1
                    while peek < len(lines) and lines[peek].strip() == '':
                        peek += 1
                    if peek < len(lines) and re.match(r'^\d+\.\s+', lines[peek].strip()):
                        i = peek
                    else:
                        break
                else:
                    break
            blocks.append(MdList(ordered=True, items=items))
            continue

        # Unordered list (with nested sub-bullet support)
        if re.match(r'^[-*]\s+', trimmed):
            items: List[MdListItem] = []
            while i < len(lines):
                raw = lines[i]
                lt = raw.strip()
                if re.match(r'^[-*]\s+', lt):
                    # Detect indent level from raw line (before stripping)
                    leading = len(raw) - len(raw.lstrip())
                    indent_level = min(leading // 2, 2)  # 0, 2+ spaces → 1, 4+ → 2
                    items.append(MdListItem(
                        text=re.sub(r'^[-*]\s+', '', lt),
                        indent=indent_level,
                    ))
                    i += 1
                elif lt == '':
                    peek = i + 1
                    while peek < len(lines) and lines[peek].strip() == '':
                        peek += 1
                    if peek < len(lines) and re.match(r'^[-*]\s+', lines[peek].strip()):
                        i = peek
                    else:
                        break
                else:
                    break
            blocks.append(MdList(ordered=False, items=items))
            continue

        # Paragraph — collect contiguous non-special lines
        para_lines: List[str] = []
        while i < len(lines):
            cur = lines[i].strip()
            if not cur:
                break
            if re.match(r'^#{1,6}\s+', cur):
                break
            if cur.startswith('!['):
                break
            if cur.startswith('>'):
                break
            if _is_table_candidate(i):
                break
            if re.match(r'^\d+\.\s+', cur):
                break
            if re.match(r'^[-*]\s+', cur):
                break
            para_lines.append(cur)
            i += 1
        blocks.append(MdParagraph(text=' '.join(para_lines)))

    return blocks


# ---------------------------------------------------------------------------
# Inline markdown formatting parser
# ---------------------------------------------------------------------------

@dataclass
class InlineRun:
    text: str
    bold: bool = False
    italic: bool = False


_INLINE_TOKEN_RE = re.compile(
    r'(`[^`]+`|\*\*\*[^*]+?\*\*\*|___[^_]+?___|\*\*[^*]+?\*\*|__[^_]+?__|'
    r'\*[^*\n]+?\*|_[^_\n]+?_)'
)


def _strip_inline_md(text: str) -> str:
    """Strip inline markdown emphasis markers."""
    t = text
    t = re.sub(r'`([^`]+)`', r'\1', t)
    t = re.sub(r'\*\*\*([^*]+)\*\*\*', r'\1', t)
    t = re.sub(r'___([^_]+)___', r'\1', t)
    t = re.sub(r'\*\*([^*]+)\*\*', r'\1', t)
    t = re.sub(r'__([^_]+)__', r'\1', t)
    t = re.sub(r'\*([^*\n]+)\*', r'\1', t)
    t = re.sub(r'_([^_\n]+)_', r'\1', t)
    return t


def _parse_inline_formatting(text: str) -> List[InlineRun]:
    """Parse inline bold/italic/code markers into typed runs."""
    tokens = _INLINE_TOKEN_RE.split(text)
    tokens = [t for t in tokens if t]
    runs: List[InlineRun] = []

    for token in tokens:
        # Bold + italic
        m = re.match(r'^\*\*\*(.+)\*\*\*$', token) or re.match(r'^___(.+)___$', token)
        if m:
            cleaned = _strip_inline_md(m.group(1))
            if cleaned:
                runs.append(InlineRun(text=cleaned, bold=True, italic=True))
            continue

        # Bold
        m = re.match(r'^\*\*(.+)\*\*$', token) or re.match(r'^__(.+)__$', token)
        if m:
            cleaned = _strip_inline_md(m.group(1))
            if cleaned:
                runs.append(InlineRun(text=cleaned, bold=True))
            continue

        # Italic
        m = re.match(r'^\*(.+)\*$', token) or re.match(r'^_(.+)_$', token)
        if m:
            cleaned = _strip_inline_md(m.group(1))
            if cleaned:
                runs.append(InlineRun(text=cleaned, italic=True))
            continue

        # Code (strip backticks, render plain)
        m = re.match(r'^`(.+)`$', token)
        if m:
            cleaned = m.group(1)
            if cleaned:
                runs.append(InlineRun(text=cleaned))
            continue

        # Plain text
        cleaned = _strip_inline_md(token)
        if cleaned:
            runs.append(InlineRun(text=cleaned))

    if not runs:
        runs.append(InlineRun(text=_strip_inline_md(text) or text))

    return runs


# ---------------------------------------------------------------------------
# Image helpers
# ---------------------------------------------------------------------------

def _decode_data_url(url: str) -> Optional[Tuple[str, bytes]]:
    """Decode a base64 data-URL into (mime_type, raw_bytes)."""
    m = re.match(r'^data:(image/[A-Za-z0-9.+-]+);base64,(.+)$', url)
    if not m:
        return None
    try:
        return m.group(1).lower(), base64.b64decode(m.group(2))
    except Exception:
        return None


def _image_dimensions_from_bytes(data: bytes, mime: str) -> Tuple[int, int]:
    """Extract pixel dimensions from raw PNG or JPEG bytes."""
    if 'png' in mime and len(data) >= 24:
        w = struct.unpack('>I', data[16:20])[0]
        h = struct.unpack('>I', data[20:24])[0]
        return w, h

    if ('jpeg' in mime or 'jpg' in mime) and len(data) >= 4:
        offset = 2
        while offset < len(data):
            if data[offset] != 0xFF:
                offset += 1
                continue
            marker = data[offset + 1]
            if marker in (0xC0, 0xC2):
                h = struct.unpack('>H', data[offset + 5:offset + 7])[0]
                w = struct.unpack('>H', data[offset + 7:offset + 9])[0]
                return w, h
            seg_len = struct.unpack('>H', data[offset + 2:offset + 4])[0]
            if seg_len == 0:
                break
            offset += 2 + seg_len

    return 500, 300


# ---------------------------------------------------------------------------
# Low-level python-docx XML helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, color: str) -> None:
    """Apply background shading to a table cell via direct XML manipulation."""
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)


def _set_cell_margins(cell, top: int, bottom: int, left: int, right: int) -> None:
    """Set cell margins in twips."""
    tc_pr = cell._element.get_or_add_tcPr()
    margins = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        margins.append(el)
    tc_pr.append(margins)


def _set_row_height(row, twips: int, rule: str = 'exact') -> None:
    """Set table row height. Rule can be 'exact' or 'atLeast'."""
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tr_height = OxmlElement('w:trHeight')
    tr_height.set(qn('w:val'), str(twips))
    tr_height.set(qn('w:hRule'), rule)
    tr_pr.append(tr_height)


def _remove_table_borders(table: DocxTable) -> None:
    """Remove all borders from a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'FFFFFF')
        borders.append(el)
    # Remove existing borders element if present
    existing = tbl_pr.find(qn('w:tblBorders'))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(borders)


def _make_bookmark(paragraph, bookmark_id: str, bookmark_name: str) -> None:
    """Wrap the existing runs of *paragraph* in a bookmark."""
    p_elem = paragraph._element
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bookmark_id))
    bm_start.set(qn('w:name'), bookmark_name)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bookmark_id))
    # Insert start before first run, end after last run
    runs = p_elem.findall(qn('w:r'))
    if runs:
        runs[0].addprevious(bm_start)
        runs[-1].addnext(bm_end)
    else:
        p_elem.append(bm_start)
        p_elem.append(bm_end)


def _make_hyperlink_paragraph(
    document: Document,
    anchor: str,
    text: str,
    font_name: str,
    font_size: Pt,
    font_color: str,
    spacing_after: int,
    line_spacing: int,
    indent_left: Optional[int] = None,
    tab_position: Optional[int] = None,
) -> Any:
    """Create a paragraph containing an internal hyperlink with a dotted tab leader."""
    paragraph = document.add_paragraph()
    p_elem = paragraph._element

    # Paragraph properties
    pPr = p_elem.get_or_add_pPr()
    # Spacing
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), str(spacing_after))
    spacing.set(qn('w:line'), str(line_spacing))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)
    # Indent
    if indent_left:
        indent = OxmlElement('w:ind')
        indent.set(qn('w:left'), str(indent_left))
        pPr.append(indent)
    # Tab stops with dot leader
    if tab_position:
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:leader'), 'dot')
        tab.set(qn('w:pos'), str(tab_position))
        tabs.append(tab)
        pPr.append(tabs)

    # Hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)

    # Text run
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))  # half-points
    rPr.append(sz)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), font_color)
    rPr.append(color)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    hyperlink.append(run)

    # Tab run
    tab_run = OxmlElement('w:r')
    tab_el = OxmlElement('w:tab')
    tab_run.append(tab_el)
    hyperlink.append(tab_run)

    p_elem.append(hyperlink)
    return paragraph


# ---------------------------------------------------------------------------
# Helper: add a styled run to a paragraph
# ---------------------------------------------------------------------------

def _set_paragraph_spacing(
    para,
    after: Optional[int] = None,
    before: Optional[int] = None,
    line: Optional[int] = None,
    line_rule: Optional[str] = None,
) -> None:
    """Set paragraph spacing via a single clean w:spacing element.

    Avoids the duplicate-element bug caused by mixing python-docx API
    (pf.space_after) with raw XML appends.  All attributes go on ONE element.
    """
    p_elem = para._element
    pPr = p_elem.get_or_add_pPr()
    # Remove any existing spacing elements
    for existing in pPr.findall(qn('w:spacing')):
        pPr.remove(existing)
    sp = OxmlElement('w:spacing')
    if after is not None:
        sp.set(qn('w:after'), str(after))
    if before is not None:
        sp.set(qn('w:before'), str(before))
    if line is not None:
        sp.set(qn('w:line'), str(line))
    if line_rule is not None:
        sp.set(qn('w:lineRule'), line_rule)
    pPr.append(sp)


def _add_run(
    paragraph,
    text: str,
    font_name: str = FONT_FAMILIES['body'],
    size: Optional[Pt] = None,
    color: Optional[str] = None,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
):
    """Add a run with explicit formatting to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.bold = bold
    run.font.italic = italic
    if underline:
        run.font.underline = True
    return run


def _add_display_run(
    paragraph,
    text: str,
    ascii_font: Optional[str] = None,
    east_font: Optional[str] = None,
    size_hp: Optional[int] = None,
    size_cs_hp: Optional[int] = None,
    color: Optional[str] = None,
    bold: bool = False,
    underline: bool = False,
):
    """Add a run with explicit font-slot control (matching reference template).

    Unlike ``_add_run`` which sets ``font.name`` (ascii slot), this helper
    lets you set ascii/hAnsi and eastAsia/cs independently — required for
    the title page where Wavetable is used for Latin text while Amsi Pro
    Narw Black is kept for eastAsia/cs slots.
    """
    run = paragraph.add_run(text)
    rPr = run._element.get_or_add_rPr()

    if ascii_font or east_font:
        rFonts = OxmlElement('w:rFonts')
        if ascii_font:
            rFonts.set(qn('w:ascii'), ascii_font)
            rFonts.set(qn('w:hAnsi'), ascii_font)
        if east_font:
            rFonts.set(qn('w:eastAsia'), east_font)
            rFonts.set(qn('w:cs'), east_font)
        rPr.append(rFonts)

    if size_hp is not None:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size_hp))
        rPr.append(sz)
    if size_cs_hp is not None:
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(size_cs_hp))
        rPr.append(szCs)

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
        bCs = OxmlElement('w:bCs')
        rPr.append(bCs)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    return run


def _hp(val: int) -> Pt:
    """Convert half-points (docx.js convention) to Pt for python-docx."""
    return Pt(val / 2)


# ---------------------------------------------------------------------------
# Main generator class
# ---------------------------------------------------------------------------

class DebriefDocxGenerator:
    """Generate a branded hackathon debrief DOCX from structured JSON content."""

    def __init__(
        self,
        content: Dict[str, Any],
        logo_dir: Optional[str] = None,
    ) -> None:
        self.content = content
        self.language: str = content.get('language', 'en')
        self.logo_dir = logo_dir
        self._bookmark_counter = 0

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def generate(self, output_path: str) -> None:
        """Build the document and save to *output_path*."""
        doc = Document()

        # Fix compatibility mode: python-docx defaults to Word 2010 (mode 14)
        # which doesn't properly honour cell-level tcMar margins.
        # Mode 15 (Word 2013+) renders them correctly.
        self._fix_compat_settings(doc)

        # Configure document-level defaults and styles to match brand
        self._configure_document_defaults(doc)
        self._configure_styles(doc)

        # ---- Section 1: Title page ----
        self._setup_title_section(doc)
        self._build_title_page_table(doc)

        # ---- Section 2: Content ----
        # CONTINUOUS break: the section-break paragraph overflows from the
        # full-bleed table to page 2, and CONTINUOUS lets the TOC start
        # immediately on that same page — no blank page in between.
        new_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
        new_section.page_width = PAGE_SIZE_A4['width']
        new_section.page_height = PAGE_SIZE_A4['height']
        new_section.top_margin = Inches(1)
        new_section.bottom_margin = Inches(1)
        new_section.left_margin = Inches(1)
        new_section.right_margin = Inches(1)

        # Empty header
        header = new_section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ''

        # Footer with page number
        footer = new_section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = fp.add_run()
        run.font.name = FONT_FAMILIES['body']
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(BRAND_COLORS['ash'])
        # Insert PAGE field
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run._element.append(fld_char_begin)
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        run._element.append(instr)
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._element.append(fld_char_end)

        # TOC title (NEW_PAGE break already starts new page — no explicit page break needed)
        self._add_toc_title(doc)

        # Static TOC
        self._add_static_toc(doc)

        # PageBreak before content
        pb_para2 = doc.add_paragraph()
        run2 = pb_para2.add_run()
        run2.add_break(WD_BREAK.PAGE)

        # Content sections
        self._add_content_sections(doc)

        # Save
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        doc.save(output_path)
        print(f'Successfully generated: {output_path}')

    # ------------------------------------------------------------------
    # Document compat settings
    # ------------------------------------------------------------------

    @staticmethod
    def _fix_compat_settings(doc: Document) -> None:
        """Set Word 2013+ compatibility mode and remove useFELayout.

        python-docx creates documents in Word 2010 mode (compatibilityMode=14)
        with useFELayout enabled.  Word 2010 mode does not honour cell-level
        tcMar margins — the fix is to switch to mode 15 (Word 2013+) and
        remove the Far East layout flag.
        """
        settings = doc.settings.element
        compat = settings.find(qn('w:compat'))
        if compat is None:
            compat = OxmlElement('w:compat')
            settings.append(compat)

        # Remove useFELayout (Far East layout — not needed, breaks margins)
        for fe in compat.findall(qn('w:useFELayout')):
            compat.remove(fe)

        # Update compatibilityMode from 14 → 15
        uri = 'http://schemas.microsoft.com/office/word'
        for cs in compat.findall(qn('w:compatSetting')):
            if (cs.get(qn('w:name')) == 'compatibilityMode'
                    and cs.get(qn('w:uri')) == uri):
                cs.set(qn('w:val'), '15')
                break
        else:
            # No existing compatibilityMode — add one
            cs = OxmlElement('w:compatSetting')
            cs.set(qn('w:name'), 'compatibilityMode')
            cs.set(qn('w:uri'), uri)
            cs.set(qn('w:val'), '15')
            compat.append(cs)

    # ------------------------------------------------------------------
    # Document defaults & styles (matches reference template)
    # ------------------------------------------------------------------

    @staticmethod
    def _configure_document_defaults(doc: Document) -> None:
        """Set w:docDefaults so every paragraph/run inherits brand fonts,
        color, size, and spacing — matching the reference template exactly."""
        styles_elem = doc.styles.element
        doc_defaults = styles_elem.find(qn('w:docDefaults'))
        if doc_defaults is None:
            doc_defaults = OxmlElement('w:docDefaults')
            styles_elem.insert(0, doc_defaults)

        # --- Run defaults ---
        rPrDefault = doc_defaults.find(qn('w:rPrDefault'))
        if rPrDefault is None:
            rPrDefault = OxmlElement('w:rPrDefault')
            doc_defaults.append(rPrDefault)
        rPr = rPrDefault.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            rPrDefault.append(rPr)

        # Font: Akkurat LL (all slots)
        for existing in rPr.findall(qn('w:rFonts')):
            rPr.remove(existing)
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), FONT_FAMILIES['body'])
        rFonts.set(qn('w:hAnsi'), FONT_FAMILIES['body'])
        rFonts.set(qn('w:eastAsia'), FONT_FAMILIES['body'])
        rFonts.set(qn('w:cs'), FONT_FAMILIES['body'])
        rPr.append(rFonts)

        # Size: 22 half-points = 11pt
        for existing in rPr.findall(qn('w:sz')):
            rPr.remove(existing)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '22')
        rPr.append(sz)
        for existing in rPr.findall(qn('w:szCs')):
            rPr.remove(existing)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '22')
        rPr.append(szCs)

        # Color: #2F2F2F
        for existing in rPr.findall(qn('w:color')):
            rPr.remove(existing)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), BRAND_COLORS['ash'])
        rPr.append(color)

        # --- Paragraph defaults ---
        pPrDefault = doc_defaults.find(qn('w:pPrDefault'))
        if pPrDefault is None:
            pPrDefault = OxmlElement('w:pPrDefault')
            doc_defaults.append(pPrDefault)
        pPr = pPrDefault.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            pPrDefault.append(pPr)

        # Spacing: after=200 line=324 lineRule=auto
        for existing in pPr.findall(qn('w:spacing')):
            pPr.remove(existing)
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:after'), '200')
        sp.set(qn('w:line'), '324')
        sp.set(qn('w:lineRule'), 'auto')
        pPr.append(sp)

    @staticmethod
    def _configure_styles(doc: Document) -> None:
        """Configure Heading 1/2/3 styles to match the reference template."""

        def _set_style_font(
            style, east_font, size_hp, color_val,
            bold=None, bold_cs=False, size_cs_hp=None,
            ascii_font=None, hAnsi_font=None,
        ):
            """Configure a style's font via XML for full control.

            By default only sets eastAsia/cs font slots — Latin text inherits
            from docDefaults (Akkurat LL).  Pass *ascii_font*/*hAnsi_font*
            to override Latin font slots (used only for Heading 3).
            """
            se = style.element
            rPr = se.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                se.append(rPr)

            # Font slots
            for existing in rPr.findall(qn('w:rFonts')):
                rPr.remove(existing)
            rFonts = OxmlElement('w:rFonts')
            if ascii_font:
                rFonts.set(qn('w:ascii'), ascii_font)
            if hAnsi_font:
                rFonts.set(qn('w:hAnsi'), hAnsi_font)
            rFonts.set(qn('w:eastAsia'), east_font)
            rFonts.set(qn('w:cs'), east_font)
            rPr.append(rFonts)

            # Size (sz + szCs)
            for existing in rPr.findall(qn('w:sz')):
                rPr.remove(existing)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(size_hp))
            rPr.append(sz)
            for existing in rPr.findall(qn('w:szCs')):
                rPr.remove(existing)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(size_cs_hp or size_hp))
            rPr.append(szCs)

            # Color (remove theme-based color, set explicit)
            for existing in rPr.findall(qn('w:color')):
                rPr.remove(existing)
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color_val)
            rPr.append(c)

            # Bold — always remove existing w:b first, then set if requested
            for existing in rPr.findall(qn('w:b')):
                rPr.remove(existing)
            if bold:
                b = OxmlElement('w:b')
                rPr.append(b)
            # Bold complex-script
            for existing in rPr.findall(qn('w:bCs')):
                rPr.remove(existing)
            if bold_cs or bold:
                bCs = OxmlElement('w:bCs')
                rPr.append(bCs)

        def _set_style_spacing(style, before=None, after=None):
            """Configure a style's paragraph spacing."""
            se = style.element
            pPr = se.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                se.append(pPr)
            for existing in pPr.findall(qn('w:spacing')):
                pPr.remove(existing)
            sp = OxmlElement('w:spacing')
            if before is not None:
                sp.set(qn('w:before'), str(before))
            if after is not None:
                sp.set(qn('w:after'), str(after))
            pPr.append(sp)

        green = BRAND_COLORS['sharpGreen']
        heading_font = FONT_FAMILIES['heading']

        # Heading 1: 16pt green, bold for all scripts.
        # eastAsia/cs = Amsi Pro Narw Black; ascii/hAnsi inherit Akkurat LL.
        h1 = doc.styles['Heading 1']
        _set_style_font(h1, heading_font, 32, green, bold=True, bold_cs=True)
        _set_style_spacing(h1, before=400)

        # Heading 2: 14pt green, bold + bCs.
        # eastAsia/cs = Amsi Pro Narw Black; ascii/hAnsi inherit Akkurat LL.
        h2 = doc.styles['Heading 2']
        _set_style_font(h2, heading_font, 28, green, bold=True)
        _set_style_spacing(h2, before=300, after=150)

        # Heading 3: 12pt green, bold + bCs, ALL font slots = Amsi Pro Narw Black.
        h3 = doc.styles['Heading 3']
        _set_style_font(
            h3, heading_font, 24, green, bold=True,
            ascii_font=heading_font, hAnsi_font=heading_font,
        )
        _set_style_spacing(h3, before=200, after=100)

        # Create bullet numbering definition for ListParagraph
        # (matches reference: Symbol font bullet, left=720, hanging=360)
        numbering_part = doc.part.numbering_part
        numbering_elem = numbering_part.element

        # abstractNum definition
        abstract_num = OxmlElement('w:abstractNum')
        abstract_num.set(qn('w:abstractNumId'), '100')
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), 'bullet')
        lvl.append(num_fmt)
        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), '\uf0b7')  # Symbol bullet
        lvl.append(lvl_text)
        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        lvl.append(lvl_jc)
        lvl_pPr = OxmlElement('w:pPr')
        lvl_ind = OxmlElement('w:ind')
        lvl_ind.set(qn('w:left'), '720')
        lvl_ind.set(qn('w:hanging'), '360')
        lvl_pPr.append(lvl_ind)
        lvl.append(lvl_pPr)
        lvl_rPr = OxmlElement('w:rPr')
        lvl_rFonts = OxmlElement('w:rFonts')
        lvl_rFonts.set(qn('w:ascii'), 'Symbol')
        lvl_rFonts.set(qn('w:hAnsi'), 'Symbol')
        lvl_rPr.append(lvl_rFonts)
        lvl.append(lvl_rPr)
        abstract_num.append(lvl)
        numbering_elem.append(abstract_num)

        # num instance pointing to abstractNum
        num = OxmlElement('w:num')
        num.set(qn('w:numId'), '100')
        abstract_ref = OxmlElement('w:abstractNumId')
        abstract_ref.set(qn('w:val'), '100')
        num.append(abstract_ref)
        numbering_elem.append(num)

    # ------------------------------------------------------------------
    # Section 1: Title page
    # ------------------------------------------------------------------

    def _setup_title_section(self, doc: Document) -> None:
        """Configure the first section for the full-bleed title page."""
        section = doc.sections[0]
        section.page_width = PAGE_SIZE_A4['width']
        section.page_height = PAGE_SIZE_A4['height']
        section.top_margin = Twips(0)
        section.bottom_margin = Twips(0)
        section.left_margin = Twips(0)
        section.right_margin = Twips(0)

    def _build_title_page_table(self, doc: Document) -> None:
        """Create the single-cell green table that forms the cover page.

        Layout and font choices match Montanstahl_Original.docx exactly:
        - Wavetable font for "hackathon"/"debrief" and "Company x One Thousand"
        - Akkurat LL (inherited from docDefaults) for subtitle, participants, date
        - Cell margins 1440 dxa (1 inch), row height 17600 twips
        """
        data = self.content.get('structuredData', self.content)
        company_name = data.get('company', {}).get('name', '')
        customer_participants = data.get('participants', {}).get('customer', [])
        ot_participants = data.get('participants', {}).get('oneThousand', [])
        total_participants = len(customer_participants) + len(ot_participants)
        profile = _get_title_page_profile(total_participants, len(company_name))
        metadata = data.get('metadata', {})
        use_cases = data.get('useCases', [])

        wavetable = FONT_FAMILIES['display_title']
        amsi = FONT_FAMILIES['heading']
        white = BRAND_COLORS['pureWhite']

        # Create table 1x1 and rebuild tblPr from scratch to match
        # the reference XML exactly (order, attributes, no jc=center).
        table = doc.add_table(rows=1, cols=1)

        tbl = table._tbl
        # Remove python-docx's auto-generated tblPr entirely
        old_pr = tbl.tblPr
        tbl.remove(old_pr)
        # Build tblPr matching reference element order exactly:
        # tblW → tblBorders → tblCellMar → tblLook
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

        tbl_w = OxmlElement('w:tblW')
        tbl_w.set(qn('w:w'), '5000')
        tbl_w.set(qn('w:type'), 'pct')
        tbl_pr.append(tbl_w)

        # Borders: all none (reference uses w:color="FFFFFF")
        tbl_borders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'), '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'FFFFFF')
            tbl_borders.append(b)
        tbl_pr.append(tbl_borders)

        tbl_cell_mar = OxmlElement('w:tblCellMar')
        for side, val in [('left', 10), ('right', 10)]:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:w'), str(val))
            el.set(qn('w:type'), 'dxa')
            tbl_cell_mar.append(el)
        tbl_pr.append(tbl_cell_mar)

        tbl_look = OxmlElement('w:tblLook')
        tbl_look.set(qn('w:val'), '0000')
        for attr in ('firstRow', 'lastRow', 'firstColumn', 'lastColumn', 'noHBand', 'noVBand'):
            tbl_look.set(qn(f'w:{attr}'), '0')
        tbl_pr.append(tbl_look)

        # gridCol: must be 11906 (full A4 width), not python-docx's 9026
        tbl_grid = tbl.find(qn('w:tblGrid'))
        if tbl_grid is not None:
            for gc in tbl_grid.findall(qn('w:gridCol')):
                gc.set(qn('w:w'), '11906')

        row = table.rows[0]
        _set_row_height(row, TABLE_ROW_HEIGHT_TWIPS)
        # cantSplit (matches reference)
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        tr_pr.append(cant_split)

        cell = row.cells[0]
        _set_cell_shading(cell, BRAND_COLORS['sharpGreen'])
        # Override cell width to auto (reference uses w:w="0" w:type="auto")
        # python-docx defaults to dxa with calculated width that's too narrow
        tc_pr = cell._element.get_or_add_tcPr()
        existing_tcW = tc_pr.find(qn('w:tcW'))
        if existing_tcW is not None:
            tc_pr.remove(existing_tcW)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), '0')
        tcW.set(qn('w:type'), 'auto')
        tc_pr.insert(0, tcW)
        # Cell margins: 1440 dxa = 1 inch (NOT Inches() which returns EMU)
        _set_cell_margins(cell, top=1440, bottom=1440, left=1440, right=1440)

        # Remove default empty paragraph from cell
        for p in cell.paragraphs:
            p_elem = p._element
            p_elem.getparent().remove(p_elem)

        # ---- Cell contents (hackathon-web TS layout + adaptive sizing) ----

        # Logo (right-aligned)
        logo_path = self._find_logo()
        if logo_path:
            p = OxmlElement('w:p')
            cell._element.append(p)
            para = cell.paragraphs[-1]
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_paragraph_spacing(para, after=profile.logoAfter)
            run = para.add_run()
            run.add_picture(logo_path, width=Pt(profile.logoSize), height=Pt(profile.logoSize))

        # Spacer before title
        self._add_cell_para(cell, '', space_before=profile.titleSpacerBefore)

        # "hackathon" — Wavetable (ascii/hAnsi), Amsi (eastAsia/cs)
        p = self._add_cell_para(cell, '', space_after=0)
        _add_display_run(p, 'hackathon',
                         ascii_font=wavetable, east_font=amsi,
                         size_hp=profile.titleSize, size_cs_hp=profile.titleSize,
                         color=white, bold=True)

        # "debrief" — same font treatment
        p = self._add_cell_para(cell, '', space_after=profile.debriefAfter)
        _add_display_run(p, 'debrief',
                         ascii_font=wavetable, east_font=amsi,
                         size_hp=profile.titleSize, size_cs_hp=profile.titleSize,
                         color=white, bold=True)

        # Use case subtitle (uppercase)
        uc_title = ''
        if use_cases:
            uc_title = use_cases[0].get('title', '')
        if not uc_title:
            uc_title = metadata.get('title', '')
        if uc_title:
            p = self._add_cell_para(cell, '', space_after=profile.subtitleAfter)
            _add_display_run(p, uc_title.upper(),
                             ascii_font=amsi, east_font=amsi,
                             size_hp=profile.subtitleSize, size_cs_hp=profile.subtitleSize,
                             color=white, bold=True)

        # Company name as section header (underlined)
        p = self._add_cell_para(cell, '', space_after=100)
        _add_display_run(p, company_name,
                         size_hp=profile.companyHeaderSize,
                         size_cs_hp=profile.companyHeaderSize,
                         color=white, underline=True)

        # Customer participants — name (role), no indent
        sorted_customer = sorted(customer_participants, key=lambda x: x.get('name', ''))
        for participant in sorted_customer:
            name = participant.get('name', '')
            role = participant.get('role', '')
            display = f'{name} ({role})' if role else name
            p = self._add_cell_para(cell, '', space_after=profile.participantAfter)
            _add_display_run(p, display,
                             size_hp=profile.participantSize,
                             size_cs_hp=profile.participantSize,
                             color=white)

        # "One Thousand" section header (underlined)
        p = self._add_cell_para(cell, '', space_before=200, space_after=100)
        _add_display_run(p, 'One Thousand',
                         size_hp=profile.companyHeaderSize,
                         size_cs_hp=profile.companyHeaderSize,
                         color=white, underline=True)

        # One Thousand participants — name (role), no indent
        sorted_ot = sorted(ot_participants, key=lambda x: x.get('name', ''))
        for participant in sorted_ot:
            name = participant.get('name', '')
            role = participant.get('role', '')
            display = f'{name} ({role})' if role else name
            p = self._add_cell_para(cell, '', space_after=profile.participantAfter)
            _add_display_run(p, display,
                             size_hp=profile.participantSize,
                             size_cs_hp=profile.participantSize,
                             color=white)

        # Footer spacer
        self._add_cell_para(cell, '', space_before=profile.footerBefore)

        # "Company x One Thousand" — Wavetable (ascii/hAnsi), Amsi (eastAsia/cs)
        p = self._add_cell_para(cell, '', space_after=100)
        _add_display_run(p, f'{company_name} x One Thousand',
                         ascii_font=wavetable, east_font=amsi,
                         size_hp=profile.footerTitleSize,
                         size_cs_hp=profile.footerTitleSize,
                         color=white, bold=True)

        # Date and location
        dates = metadata.get('dates')
        if dates:
            start_str = self._format_date_ddmmyyyy(dates.get('start', ''))
            end_str = self._format_date_ddmmyyyy(dates.get('end', ''))
            date_str = f'{start_str} - {end_str}'
        else:
            date_str = metadata.get('date', '')
        location = metadata.get('location', '')
        location_date = f'{location}, {date_str}' if location else date_str

        p = self._add_cell_para(cell, '')
        _add_display_run(p, location_date,
                         size_hp=profile.dateSize, size_cs_hp=profile.dateSize,
                         color=white, bold=True)

    def _add_cell_para(
        self,
        cell,
        text: str,
        space_before: int = 0,
        space_after: int = 0,
    ):
        """Add a paragraph to a table cell with spacing (in twips).

        Only sets ``space_before`` / ``space_after`` when non-zero to keep the
        XML clean (Word treats omitted values as 0).
        """
        p_elem = OxmlElement('w:p')
        cell._element.append(p_elem)
        para = cell.paragraphs[-1]
        # Use raw XML; always emit 'after' to override docDefault after=200.
        # Only emit 'before' when non-zero (matches reference pattern).
        pPr = para._element.get_or_add_pPr()
        sp = OxmlElement('w:spacing')
        if space_before:
            sp.set(qn('w:before'), str(space_before))
        sp.set(qn('w:after'), str(space_after))
        pPr.append(sp)
        if text:
            _add_run(para, text)
        return para

    @staticmethod
    def _set_para_indent(para, left: int = 0) -> None:
        """Set paragraph left indent in twips (dxa)."""
        p_elem = para._element
        pPr = p_elem.get_or_add_pPr()
        # Remove existing indent
        for existing in pPr.findall(qn('w:ind')):
            pPr.remove(existing)
        if left:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), str(left))
            pPr.append(ind)

    # ------------------------------------------------------------------
    # Section 2: TOC
    # ------------------------------------------------------------------

    def _add_toc_title(self, doc: Document) -> None:
        """Add the TOC heading.

        Reference uses eastAsia=Amsi Pro Narw Black only (ascii/hAnsi inherit
        Akkurat LL from docDefaults).  sz=56 (28pt), green, bold.
        """
        title = 'INHALTSVERZEICHNIS' if self.language == 'de' else 'TABLE OF CONTENTS'
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.space_after = Twips(600)
        _add_display_run(para, title,
                         east_font=FONT_FAMILIES['heading'],
                         size_hp=56, size_cs_hp=56,
                         color=BRAND_COLORS['sharpGreen'], bold=True)

    def _add_static_toc(self, doc: Document) -> None:
        """Build clickable TOC entries with dotted tab leaders."""
        data = self.content.get('structuredData', self.content)
        sections = data.get('sections', self.content.get('sections', []))
        use_cases = data.get('useCases', [])
        entries = _build_toc_entries(sections, self.language, use_cases)

        tab_pos = int(Inches(6))  # right-aligned tab stop

        for entry in entries:
            bookmark_name = f'section-{entry.section_id}' if entry.section_id else None
            # Build display text: "1. Title" for level 1, "1.1 Title" for level 2
            if entry.level == 1:
                list_text = f'{entry.number_label}. {entry.title}'
            else:
                list_text = f'{entry.number_label} {entry.title}'

            font_size = Pt(11) if entry.level == 1 else Pt(10)
            indent_left = int(Inches(0.28)) if entry.level == 2 else None
            sp_after = 160 if entry.level == 1 else 120

            if bookmark_name:
                _make_hyperlink_paragraph(
                    doc, bookmark_name, list_text,
                    FONT_FAMILIES['body'], font_size, BRAND_COLORS['ash'],
                    spacing_after=sp_after, line_spacing=280,
                    indent_left=indent_left, tab_position=tab_pos,
                )
            else:
                para = doc.add_paragraph()
                pf = para.paragraph_format
                pf.space_after = Twips(sp_after)
                # Set line spacing
                p_elem = para._element
                pPr = p_elem.get_or_add_pPr()
                sp = OxmlElement('w:spacing')
                sp.set(qn('w:line'), '280')
                sp.set(qn('w:lineRule'), 'auto')
                pPr.append(sp)
                if indent_left:
                    ind = OxmlElement('w:ind')
                    ind.set(qn('w:left'), str(indent_left))
                    pPr.append(ind)
                _add_run(para, list_text, FONT_FAMILIES['body'], font_size,
                         BRAND_COLORS['ash'])

    # ------------------------------------------------------------------
    # Section 2: Content sections
    # ------------------------------------------------------------------

    def _add_content_sections(self, doc: Document) -> None:
        """Add all content sections in canonical order."""
        data = self.content.get('structuredData', self.content)
        sections = data.get('sections', self.content.get('sections', []))
        by_id: Dict[str, Dict[str, Any]] = {s['id']: s for s in sections}

        sorted_ids = [sid for sid in SECTION_ORDER if sid in by_id]

        for section_id in sorted_ids:
            section = by_id[section_id]
            bookmark_name = f'section-{section_id}'

            # H1 heading with bookmark — uses Heading1 style (no explicit run formatting)
            para = doc.add_paragraph(style='Heading 1')
            run = para.add_run(section.get('title', section_id))
            # Style handles font/size/color; only set spacing override
            _set_paragraph_spacing(para, before=400, after=300)

            self._bookmark_counter += 1
            _make_bookmark(para, self._bookmark_counter, bookmark_name)

            # Convert markdown content
            content_text = section.get('content', '')
            title = section.get('title', '')
            cleaned = _strip_code_blocks(_strip_redundant_heading(title, content_text))
            self._convert_markdown(doc, cleaned)

    def _convert_markdown(self, doc: Document, markdown: str) -> None:
        """Convert markdown text into document paragraphs and tables."""
        blocks = _parse_markdown_blocks(markdown)

        for block in blocks:
            if isinstance(block, MdHeading):
                level = block.level
                style_name = 'Heading 2' if level <= 2 else 'Heading 3'

                # Use style — no explicit font/size/color on runs
                para = doc.add_paragraph(style=style_name)
                para.add_run(_strip_inline_md(block.text))

            elif isinstance(block, MdList):
                for idx, item in enumerate(block.items):
                    # Extract text and indent level
                    if isinstance(item, MdListItem):
                        item_text = item.text
                        indent_level = item.indent
                    else:
                        item_text = str(item)
                        indent_level = 0

                    if block.ordered:
                        list_text = f'{idx + 1}. {item_text}'
                    else:
                        list_text = item_text

                    if block.ordered:
                        # Ordered items: regular paragraphs with number prefix
                        # (inherits after=200 + line=324 from docDefaults)
                        para = doc.add_paragraph()
                    else:
                        # Bullet list via List Paragraph style (matches reference)
                        para = doc.add_paragraph(style='List Paragraph')
                        _set_paragraph_spacing(para, after=120)

                        # Connect to bullet numbering definition (numId=100)
                        p_elem = para._element
                        pPr = p_elem.get_or_add_pPr()
                        numPr = OxmlElement('w:numPr')
                        ilvl = OxmlElement('w:ilvl')
                        ilvl.set(qn('w:val'), str(indent_level))
                        numPr.append(ilvl)
                        numId = OxmlElement('w:numId')
                        numId.set(qn('w:val'), '100')
                        numPr.append(numId)
                        pPr.append(numPr)

                        # Apply additional indent for nested sub-bullets
                        if indent_level > 0:
                            ind = OxmlElement('w:ind')
                            indent_twips = str(720 + indent_level * 360)
                            ind.set(qn('w:left'), indent_twips)
                            ind.set(qn('w:hanging'), '360')
                            pPr.append(ind)

                    self._add_inline_runs(para, list_text)

            elif isinstance(block, MdImage):
                self._add_image_block(doc, block.url, block.alt)

            elif isinstance(block, MdTable):
                self._add_table_block(doc, block)

            elif isinstance(block, MdParagraph):
                if not block.text.strip():
                    continue
                # Normal paragraph — inherits after=200, line=324 from docDefaults
                para = doc.add_paragraph()
                self._add_inline_runs(para, block.text)

    def _add_inline_runs(self, paragraph, text: str) -> None:
        """Parse inline formatting and add runs to *paragraph*.

        Runs inherit font/size/color from document defaults and styles.
        Only bold/italic are set explicitly when needed.
        """
        runs = _parse_inline_formatting(text)
        for r in runs:
            run = paragraph.add_run(r.text)
            if r.bold:
                run.bold = True
            if r.italic:
                run.italic = True

    def _add_image_block(self, doc: Document, url: str, alt: str) -> None:
        """Embed a base64 data-URL image or add a placeholder."""
        decoded = _decode_data_url(url)
        if not decoded:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = para.paragraph_format
            pf.space_after = Twips(220)
            label = f'[Image: {alt}]' if alt else '[Image]'
            _add_run(para, label, FONT_FAMILIES['body'], Pt(11),
                     BRAND_COLORS['ceruleanBlue'], italic=True)
            return

        mime, img_bytes = decoded
        w, h = _image_dimensions_from_bytes(img_bytes, mime)
        max_w, max_h = 550, 320
        scale_w = max_w / w if w > 0 else 1
        scale_h = max_h / h if h > 0 else 1
        scale = min(scale_w, scale_h, 1)
        final_w = max(120, round(w * scale))
        final_h = max(90, round(h * scale))

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = para.paragraph_format
        pf.space_before = Twips(200)
        pf.space_after = Twips(80)
        run = para.add_run()
        image_stream = io.BytesIO(img_bytes)
        run.add_picture(image_stream, width=Pt(final_w), height=Pt(final_h))

        # Caption
        if alt and alt.strip():
            cap_para = doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_pf = cap_para.paragraph_format
            cap_pf.space_after = Twips(220)
            _add_run(cap_para, alt.strip(), FONT_FAMILIES['body'], Pt(9),
                     BRAND_COLORS['ceruleanBlue'], italic=True)

    def _add_table_block(self, doc: Document, block: MdTable) -> None:
        """Render a markdown table with header shading and thin borders."""
        col_count = max(len(block.headers), 1)

        # Spacer before
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Twips(200)

        # Normalise rows to have the correct column count
        def _norm(row: List[str]) -> List[str]:
            if len(row) < col_count:
                return row + [''] * (col_count - len(row))
            return row[:col_count]

        total_rows = 1 + len(block.rows)
        table = doc.add_table(rows=total_rows, cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set thin gray borders
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        borders = OxmlElement('w:tblBorders')
        for side in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), 'D1D5DB')
            borders.append(el)
        existing = tbl_pr.find(qn('w:tblBorders'))
        if existing is not None:
            tbl_pr.remove(existing)
        tbl_pr.append(borders)

        # Full-width table
        tbl_w = OxmlElement('w:tblW')
        tbl_w.set(qn('w:w'), '5000')
        tbl_w.set(qn('w:type'), 'pct')
        existing_w = tbl_pr.find(qn('w:tblW'))
        if existing_w is not None:
            tbl_pr.remove(existing_w)
        tbl_pr.append(tbl_w)

        # Header row
        headers_norm = _norm(block.headers)
        for ci, header_text in enumerate(headers_norm):
            cell = table.rows[0].cells[ci]
            _set_cell_shading(cell, 'F5F5F5')
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Twips(0)
            self._add_inline_runs(p, header_text)

        # Data rows
        for ri, row_data in enumerate(block.rows):
            norm_row = _norm(row_data)
            for ci, cell_text in enumerate(norm_row):
                cell = table.rows[ri + 1].cells[ci]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Twips(0)
                self._add_inline_runs(p, cell_text)

        # Spacer after
        sp2 = doc.add_paragraph()
        sp2.paragraph_format.space_after = Twips(200)

    # ------------------------------------------------------------------
    # Utility
    # ------------------------------------------------------------------

    @staticmethod
    def _format_date_ddmmyyyy(date_str: str) -> str:
        """Convert ISO date (YYYY-MM-DD) to DD.MM.YYYY. Pass through other formats."""
        m = re.match(r'^(\d{4})-(\d{2})-(\d{2})$', date_str.strip())
        if m:
            return f'{m.group(3)}.{m.group(2)}.{m.group(1)}'
        return date_str

    def _find_logo(self) -> Optional[str]:
        """Locate the logo PNG file."""
        filename = 'onethousand-icon-limeonblack-rounded.png'
        if self.logo_dir:
            candidate = os.path.join(self.logo_dir, filename)
            if os.path.isfile(candidate):
                return candidate
        return None


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> int:
    """Entry point for CLI invocation."""
    parser = argparse.ArgumentParser(
        description='Generate branded hackathon debrief DOCX from JSON content',
    )
    parser.add_argument(
        '--content',
        required=True,
        help='Path to the JSON content file',
    )
    parser.add_argument(
        '--logo-dir',
        default=None,
        help='Path to the directory containing logo PNG files',
    )
    parser.add_argument(
        '--output',
        required=True,
        help='Output DOCX file path',
    )

    args = parser.parse_args()

    # Validate inputs
    if not os.path.isfile(args.content):
        print(f'Error: Content file not found: {args.content}', file=sys.stderr)
        return 1

    # Load JSON
    try:
        with open(args.content, 'r', encoding='utf-8') as f:
            content = json.load(f)
        print(f'Loaded content from {args.content}')
    except json.JSONDecodeError as exc:
        print(f'Error parsing JSON: {exc}', file=sys.stderr)
        return 1

    if args.logo_dir and not os.path.isdir(args.logo_dir):
        print(f'Warning: Logo directory not found: {args.logo_dir}', file=sys.stderr)

    # Generate
    try:
        generator = DebriefDocxGenerator(content, logo_dir=args.logo_dir)
        generator.generate(args.output)
        return 0
    except Exception as exc:
        print(f'Error generating document: {exc}', file=sys.stderr)
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())

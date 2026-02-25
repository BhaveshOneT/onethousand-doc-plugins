#!/usr/bin/env python3
"""
Generate scope documents from a DOCX template.

This script takes a pre-unpacked DOCX template directory, injects content via XML
placeholder replacement, and repacks it into a final .docx file.

Usage:
    python generate_scope_doc.py --template-dir <path> --variables <vars.json> \
        --content <content.json> --output <output.docx> [--arch-diagram <image_path>]
"""

import argparse
import json
import shutil
import tempfile
import zipfile
import os
import re
import random
import sys
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple

# Try to import PIL for image sizing, but handle gracefully if not available
try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False


def escape_xml_text(text: str) -> str:
    """
    Escape special XML characters in text content.

    Args:
        text: Plain text to escape

    Returns:
        XML-safe text with &, <, > escaped
    """
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    return text


def generate_paragraph_id() -> str:
    """
    Generate a unique paragraph ID (random hex string).

    Returns:
        A random hex string suitable for w14:paraId
    """
    return format(random.randint(0, 0xFFFFFFFF), '08x')


def get_image_dimensions(image_path: str, max_width_inches: float = 6.0) -> Tuple[int, int]:
    """
    Get image dimensions in EMUs (English Metric Units), capped to page width.

    Word uses EMUs where 914400 EMUs = 1 inch.
    The image is scaled down proportionally if it exceeds *max_width_inches*.
    Default fallback: 6 × 4 inches.

    Args:
        image_path: Path to the image file
        max_width_inches: Maximum width in inches (default 6.0 for letter with margins)

    Returns:
        Tuple of (width_emu, height_emu)
    """
    default_width_emu = 5486400  # ~6 inches
    default_height_emu = 3657600  # ~4 inches
    emu_per_inch = 914400

    if not HAS_PIL:
        return default_width_emu, default_height_emu

    try:
        with Image.open(image_path) as img:
            width_px, height_px = img.size
            # Try to read actual DPI from the image; fall back to 96
            dpi_info = img.info.get('dpi', (96, 96))
            dpi_x = dpi_info[0] if dpi_info[0] > 0 else 96

            width_inches = width_px / dpi_x
            height_inches = height_px / dpi_x

            # Scale down if wider than max_width_inches
            if width_inches > max_width_inches:
                scale = max_width_inches / width_inches
                width_inches = max_width_inches
                height_inches *= scale

            width_emu = int(width_inches * emu_per_inch)
            height_emu = int(height_inches * emu_per_inch)

            return width_emu, height_emu
    except Exception as e:
        print(f"Warning: Could not determine image dimensions: {e}")
        return default_width_emu, default_height_emu


def generate_image_xml(image_path: str, rel_id: str) -> str:
    """
    Generate OOXML for an inline image.

    Args:
        image_path: Path to the image file
        rel_id: Relationship ID to use in r:embed

    Returns:
        OOXML string for the image
    """
    width_emu, height_emu = get_image_dimensions(image_path)

    image_xml = (
        f'<w:p>'
        f'\n  <w:pPr><w:jc w:val="center"/></w:pPr>'
        f'\n  <w:r>'
        f'\n    <w:drawing>'
        f'\n      <wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'\n        <wp:extent cx="{width_emu}" cy="{height_emu}"/>'
        f'\n        <wp:docPr id="100" name="Architecture Diagram"/>'
        f'\n        <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        f'\n          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'\n            <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'\n              <pic:nvPicPr>'
        f'\n                <pic:cNvPr id="100" name="arch_diagram"/>'
        f'\n                <pic:cNvPicPr/>'
        f'\n              </pic:nvPicPr>'
        f'\n              <pic:blipFill>'
        f'\n                <a:blip r:embed="{rel_id}"/>'
        f'\n                <a:stretch><a:fillRect/></a:stretch>'
        f'\n              </pic:blipFill>'
        f'\n              <pic:spPr>'
        f'\n                <a:xfrm>'
        f'\n                  <a:off x="0" y="0"/>'
        f'\n                  <a:ext cx="{width_emu}" cy="{height_emu}"/>'
        f'\n                </a:xfrm>'
        f'\n                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'\n              </pic:spPr>'
        f'\n            </pic:pic>'
        f'\n          </a:graphicData>'
        f'\n        </a:graphic>'
        f'\n      </wp:inline>'
        f'\n    </w:drawing>'
        f'\n  </w:r>'
        f'\n</w:p>'
    )

    return image_xml


def generate_single_paragraph_xml(
    text: str,
    style: str,
    lang_val: str = "en-US",
    bold: bool = False,
    numPr: Optional[str] = None,
) -> str:
    """
    Build a single <w:p> element for a paragraph.

    Args:
        text: Plain text content
        style: Paragraph style name
        lang_val: Language code (e.g., "en-US")
        bold: Whether to make text bold
        numPr: Optional numeric properties for lists (XML string)

    Returns:
        OOXML paragraph string
    """
    para_id = generate_paragraph_id()
    text_id = generate_paragraph_id()

    # Check if text has leading/trailing whitespace
    preserve_attr = ' xml:space="preserve"' if (text and (text[0] == ' ' or text[-1] == ' ')) else ''

    # Build paragraph properties
    ppr_parts = [f'<w:pStyle w:val="{style}"/>']
    if numPr:
        ppr_parts.append(numPr)

    ppr = '\n    '.join(ppr_parts)

    # Build run properties
    rpr_parts = []
    if bold:
        rpr_parts.append('<w:b/>')
    rpr_parts.append(f'<w:lang w:val="{lang_val}"/>')
    rpr = '\n    '.join(rpr_parts)

    return (
        f'<w:p w14:paraId="{para_id}" w14:textId="{text_id}">'
        f'\n  <w:pPr>'
        f'\n    {ppr}'
        f'\n  </w:pPr>'
        f'\n  <w:r>'
        f'\n    <w:rPr>'
        f'\n      {rpr}'
        f'\n    </w:rPr>'
        f'\n    <w:t{preserve_attr}>{text}</w:t>'
        f'\n  </w:r>'
        f'\n</w:p>'
    )


def generate_paragraph_xml(
    text: str,
    style: str = "Normal",
    language: str = "en",
    bold: bool = False
) -> str:
    """
    Generate OOXML paragraph(s) with text.

    If the text contains double-newlines (paragraph breaks), each block
    becomes its own <w:p> element. The first block uses the requested
    style; subsequent blocks use "Normal".

    Args:
        text: Text content
        style: Paragraph style name
        language: Language code ("en" or "de")
        bold: Whether to make text bold

    Returns:
        OOXML string for one or more paragraphs
    """
    lang_val = "en-US" if language == "en" else "de-DE"

    # Split on paragraph breaks (double newline)
    blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
    if not blocks:
        blocks = [text]

    parts = []
    for i, block in enumerate(blocks):
        # Escape XML special characters
        safe_text = escape_xml_text(block)

        # Use requested style for first block, Normal for continuations
        block_style = style if i == 0 else "Normal"

        parts.append(generate_single_paragraph_xml(safe_text, block_style, lang_val, bold))

    return "\n".join(parts)


def generate_bullet_paragraph_xml(
    text: str,
    language: str = "en",
    is_dash: bool = False
) -> str:
    """
    Generate a bullet list paragraph.

    Args:
        text: Text content for the bullet
        language: Language code
        is_dash: If True, use em-dash prefix instead of bullet

    Returns:
        OOXML string for the bullet paragraph
    """
    lang_val = "en-US" if language == "en" else "de-DE"
    para_id = generate_paragraph_id()
    text_id = generate_paragraph_id()

    # Escape XML special characters
    safe_text = escape_xml_text(text)

    if is_dash:
        # Dash/hyphen bullets: use hyphen prefix "-  " with ListParagraph indent,
        # but NO numPr (no bullet character). The hyphen IS the visual marker.
        # This matches the RVT original format for out-of-scope items.
        safe_text = "-  " + safe_text

        bullet_xml = (
            f'<w:p w14:paraId="{para_id}" w14:textId="{text_id}">'
            f'\n  <w:pPr>'
            f'\n    <w:pStyle w:val="ListParagraph"/>'
            f'\n  </w:pPr>'
            f'\n  <w:r>'
            f'\n    <w:rPr>'
            f'\n      <w:lang w:val="{lang_val}"/>'
            f'\n    </w:rPr>'
            f'\n    <w:t xml:space="preserve">{safe_text}</w:t>'
            f'\n  </w:r>'
            f'\n</w:p>'
        )
    else:
        # Normal bullets: use numPr with round bullet
        bullet_xml = (
            f'<w:p w14:paraId="{para_id}" w14:textId="{text_id}">'
            f'\n  <w:pPr>'
            f'\n    <w:pStyle w:val="ListParagraph"/>'
            f'\n    <w:numPr>'
            f'\n      <w:ilvl w:val="0"/>'
            f'\n      <w:numId w:val="27"/>'
            f'\n    </w:numPr>'
            f'\n  </w:pPr>'
            f'\n  <w:r>'
            f'\n    <w:rPr>'
            f'\n      <w:lang w:val="{lang_val}"/>'
            f'\n    </w:rPr>'
            f'\n    <w:t>{safe_text}</w:t>'
            f'\n  </w:r>'
            f'\n</w:p>'
        )

    return bullet_xml


def generate_sprint_paragraph_xml(
    sprint_label: str,
    sprint_description: str,
    language: str = "en"
) -> str:
    """
    Generate a sprint design paragraph as a bold Normal paragraph.

    In the RVT format, sprint labels are bold Normal paragraphs (not
    list items), followed by separate bullet lists for deliverables.
    Format: "Sprint 0: Title text" all in bold.

    Args:
        sprint_label: Label (e.g., "Sprint 1:")
        sprint_description: Description text
        language: Language code

    Returns:
        OOXML paragraph with fully bold sprint text
    """
    lang_val = "en-US" if language == "en" else "de-DE"
    para_id = generate_paragraph_id()
    text_id = generate_paragraph_id()

    safe_label = escape_xml_text(sprint_label)
    safe_desc = escape_xml_text(sprint_description)

    # Combine label and description into a single bold text
    full_text = f"{safe_label} {safe_desc}" if safe_desc else safe_label

    sprint_xml = (
        f'<w:p w14:paraId="{para_id}" w14:textId="{text_id}">'
        f'\n  <w:pPr>'
        f'\n    <w:pStyle w:val="Normal"/>'
        f'\n  </w:pPr>'
        f'\n  <w:r>'
        f'\n    <w:rPr>'
        f'\n      <w:b/>'
        f'\n      <w:lang w:val="{lang_val}"/>'
        f'\n    </w:rPr>'
        f'\n    <w:t xml:space="preserve">{full_text}</w:t>'
        f'\n  </w:r>'
        f'\n</w:p>'
    )

    return sprint_xml


def generate_heading3_xml(
    text: str,
    language: str = "en"
) -> str:
    """
    Generate a Heading3 paragraph (green, bold, indented sub-label).

    Args:
        text: Label text
        language: Language code

    Returns:
        OOXML paragraph with Heading3 style
    """
    lang_val = "en-US" if language == "en" else "de-DE"
    para_id = generate_paragraph_id()
    text_id = generate_paragraph_id()

    safe_text = escape_xml_text(text)

    heading3_xml = (
        f'<w:p w14:paraId="{para_id}" w14:textId="{text_id}">'
        f'\n  <w:pPr>'
        f'\n    <w:pStyle w:val="Heading3"/>'
        f'\n  </w:pPr>'
        f'\n  <w:r>'
        f'\n    <w:rPr>'
        f'\n      <w:lang w:val="{lang_val}"/>'
        f'\n    </w:rPr>'
        f'\n    <w:t>{safe_text}</w:t>'
        f'\n  </w:r>'
        f'\n</w:p>'
    )

    return heading3_xml


def is_sprint_section(section: Dict[str, Any]) -> bool:
    """
    Check if a section is a sprint design section.

    Args:
        section: Section dictionary

    Returns:
        True if section is a sprint section
    """
    section_num = section.get("number", "")
    section_title = section.get("title", "")
    is_sprint_flag = section.get("is_sprint_section", False)

    if is_sprint_flag:
        return True

    if str(section_num).startswith("Sprint"):
        return True

    if "Sprint" in section_title or "sprint" in section_title.lower():
        return True

    return False


def generate_body_content_xml(
    content_data: Dict[str, Any],
    language: str = "en"
) -> str:
    """
    Generate OOXML body content from content.json sections.

    Args:
        content_data: Parsed content.json with sections
        language: Language code for text

    Returns:
        OOXML string for all body content
    """
    xml_parts = []

    # Page break before first section
    xml_parts.append('<w:p><w:pPr><w:pageBreakBefore/></w:pPr></w:p>')

    sections = content_data.get("sections", [])

    for section_idx, section in enumerate(sections):
        # Add space between major sections (except first)
        if section_idx > 0:
            para_id = generate_paragraph_id()
            text_id = generate_paragraph_id()
            xml_parts.append(
                f'<w:p w14:paraId="{para_id}" w14:textId="{text_id}">'
                f'<w:pPr></w:pPr></w:p>'
            )

        # Section heading (Heading1)
        section_num = section.get("number", "")
        section_title = section.get("title", "")

        # Format: "1.  Initial Context" (number with dot, two spaces, title)
        if section_num:
            heading_text = f"{section_num}.  {section_title}"
        else:
            heading_text = section_title

        xml_parts.append(generate_paragraph_xml(
            heading_text,
            style="Heading1",
            language=language,
            bold=False
        ))

        # Section intro content
        section_content = section.get("content", "").strip()
        if section_content:
            xml_parts.append(generate_paragraph_xml(
                section_content,
                style="Normal",
                language=language
            ))

        # Section-level bullet points (before subsections)
        section_bullets = section.get("bullet_points", [])
        bullet_style = section.get("bullet_style", "normal")
        is_dash_bullets = bullet_style == "dash"

        for bullet in section_bullets:
            xml_parts.append(generate_bullet_paragraph_xml(
                bullet,
                language=language,
                is_dash=is_dash_bullets
            ))

        # Check if this is a sprint section
        if is_sprint_section(section):
            # Sprint subsections use bold inline labels with ListParagraph style
            subsections = section.get("subsections", [])
            for subsection in subsections:
                subsection_num = subsection.get("number", "")
                subsection_title = subsection.get("title", "")

                # Format sprint label: "Sprint 0: " or similar
                sprint_label = f"{subsection_num}:"
                sprint_desc = subsection_title

                xml_parts.append(generate_sprint_paragraph_xml(
                    sprint_label,
                    sprint_desc,
                    language=language
                ))

                subsection_content = subsection.get("content", "").strip()
                if subsection_content:
                    # Label-like content (e.g., "Deliverables:") should be bold
                    is_label = subsection_content.endswith(":") and len(subsection_content) < 50
                    xml_parts.append(generate_paragraph_xml(
                        subsection_content,
                        style="Normal",
                        language=language,
                        bold=is_label
                    ))

                # Bullet points under sprint (deliverables)
                sprint_bullets = subsection.get("bullet_points", [])
                for bullet in sprint_bullets:
                    xml_parts.append(generate_bullet_paragraph_xml(
                        bullet,
                        language=language,
                        is_dash=False
                    ))
        else:
            # Normal subsections use Heading2 with manual numbering override
            subsections = section.get("subsections", [])
            for subsection in subsections:
                subsection_num = subsection.get("number", "")
                subsection_title = subsection.get("title", "")

                # Format: "2.1    Email classification"
                if subsection_num:
                    subsection_heading = f"{subsection_num}    {subsection_title}"
                else:
                    subsection_heading = subsection_title

                # For Heading2, disable auto-numbering and use manual numbering
                numPr_override = (
                    '<w:numPr>'
                    '<w:ilvl w:val="0"/>'
                    '<w:numId w:val="0"/>'
                    '</w:numPr>'
                )

                lang_val = "en-US" if language == "en" else "de-DE"
                para_id = generate_paragraph_id()
                text_id = generate_paragraph_id()
                safe_text = escape_xml_text(subsection_heading)

                heading2_xml = (
                    f'<w:p w14:paraId="{para_id}" w14:textId="{text_id}">'
                    f'\n  <w:pPr>'
                    f'\n    <w:pStyle w:val="Heading2"/>'
                    f'\n    {numPr_override}'
                    f'\n  </w:pPr>'
                    f'\n  <w:r>'
                    f'\n    <w:rPr>'
                    f'\n      <w:lang w:val="{lang_val}"/>'
                    f'\n    </w:rPr>'
                    f'\n    <w:t xml:space="preserve">{safe_text}</w:t>'
                    f'\n  </w:r>'
                    f'\n</w:p>'
                )

                xml_parts.append(heading2_xml)

                subsection_content = subsection.get("content", "").strip()
                if subsection_content:
                    xml_parts.append(generate_paragraph_xml(
                        subsection_content,
                        style="Normal",
                        language=language
                    ))

                # Bullet points under subsection
                subsection_bullets = subsection.get("bullet_points", [])
                for bullet in subsection_bullets:
                    xml_parts.append(generate_bullet_paragraph_xml(
                        bullet,
                        language=language,
                        is_dash=False
                    ))

    return "\n".join(xml_parts)


def replace_cover_placeholders(
    document_xml: str,
    variables: Dict[str, str]
) -> str:
    """
    Replace cover page placeholders in document.xml.

    Handles newline conversion to <w:br/> for multi-line cover text.

    Args:
        document_xml: Content of document.xml
        variables: Dictionary of variables for replacement

    Returns:
        Modified document.xml content
    """
    # Map of placeholder to variable key
    cover_placeholders = {
        "{{COVER_TITLE}}": "cover_title",
        "{{COVER_SUBTITLE}}": "cover_subtitle",
        "{{COVER_CLIENT_X_OT}}": "cover_client_x_ot",
    }

    for placeholder, var_key in cover_placeholders.items():
        value = variables.get(var_key, "")

        # Handle newlines in cover text by converting to Word break elements.
        # Since the placeholder is inside <w:t>{{PLACEHOLDER}}</w:t>,
        # we replace the text with: LINE1</w:t><w:br/><w:t>LINE2
        # This keeps the outer <w:t> tags and inserts a <w:br/> between lines.
        if "\n" in value:
            lines = value.split("\n")
            escaped_lines = [escape_xml_text(line) for line in lines]
            replacement = "</w:t><w:br/><w:t>".join(escaped_lines)
        else:
            replacement = escape_xml_text(value)

        document_xml = document_xml.replace(placeholder, replacement)

    return document_xml


def add_image_to_docx(
    working_dir: str,
    image_path: str,
    image_filename: str = "image2.jpeg"
) -> str:
    """
    Add an image to the DOCX template and update relationships.

    Args:
        working_dir: Path to unpacked DOCX template
        image_path: Path to the image file to add
        image_filename: Desired filename in word/media/

    Returns:
        Relationship ID for the image
    """
    # Create media directory if it doesn't exist
    media_dir = os.path.join(working_dir, "word", "media")
    os.makedirs(media_dir, exist_ok=True)

    # Copy image to media directory
    dest_path = os.path.join(media_dir, image_filename)
    shutil.copy2(image_path, dest_path)
    print(f"Copied image to {dest_path}")

    # Update document.xml.rels to add relationship
    rels_path = os.path.join(working_dir, "word", "_rels", "document.xml.rels")

    if not os.path.exists(rels_path):
        print(f"Warning: {rels_path} not found")
        return "rIdArch"

    with open(rels_path, 'r', encoding='utf-8') as f:
        rels_content = f.read()

    # Find the next available rId number
    rel_ids = re.findall(r'Id="rId(\d+)"', rels_content)
    next_id = max([int(rid) for rid in rel_ids] + [0]) + 1
    rel_id = f"rId{next_id}"

    # Determine media type based on file extension
    ext = os.path.splitext(image_path)[1].lower()
    media_type_map = {
        '.jpeg': 'image/jpeg',
        '.jpg': 'image/jpeg',
        '.png': 'image/png',
        '.gif': 'image/gif',
        '.bmp': 'image/bmp'
    }
    media_type = media_type_map.get(ext, 'image/jpeg')

    # Add relationship entry before the closing </Relationships>
    new_rel = f'  <Relationship Id="{rel_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/{image_filename}"/>\n'
    rels_content = rels_content.replace('</Relationships>', f'{new_rel}</Relationships>')

    with open(rels_path, 'w', encoding='utf-8') as f:
        f.write(rels_content)
    print(f"Added relationship {rel_id} for image")

    # Update [Content_Types].xml if needed
    content_types_path = os.path.join(working_dir, "[Content_Types].xml")
    if os.path.exists(content_types_path):
        with open(content_types_path, 'r', encoding='utf-8') as f:
            content_types = f.read()

        # Check if this specific media file override already exists
        if f'/word/media/{image_filename}' not in content_types:
            override_entry = f'  <Override PartName="/word/media/{image_filename}" ContentType="{media_type}"/>\n'
            content_types = content_types.replace('</Types>', f'{override_entry}</Types>')

            with open(content_types_path, 'w', encoding='utf-8') as f:
                f.write(content_types)
            print(f"Added content type for {media_type}")

    return rel_id


def process_template(
    template_dir: str,
    variables: Dict[str, str],
    content_data: Dict[str, Any],
    output_path: str,
    arch_diagram_path: Optional[str] = None
) -> bool:
    """
    Process the template and generate the final DOCX.

    Args:
        template_dir: Path to unpacked DOCX template
        variables: Variables for placeholder replacement
        content_data: Content data with sections
        output_path: Output DOCX file path
        arch_diagram_path: Optional path to architecture diagram image

    Returns:
        True if successful, False otherwise
    """
    try:
        # Create a temporary working directory
        with tempfile.TemporaryDirectory() as work_temp:
            print(f"Creating working copy of template...")

            # Copy template to working directory
            working_dir = os.path.join(work_temp, "docx_work")
            shutil.copytree(template_dir, working_dir)

            # Restore [Content_Types].xml if renamed for zip compatibility
            safe_ct = os.path.join(working_dir, "_Content_Types_.xml")
            real_ct = os.path.join(working_dir, "[Content_Types].xml")
            if os.path.exists(safe_ct) and not os.path.exists(real_ct):
                os.rename(safe_ct, real_ct)

            # Read document.xml
            doc_xml_path = os.path.join(working_dir, "word", "document.xml")
            if not os.path.exists(doc_xml_path):
                print(f"Error: {doc_xml_path} not found")
                return False

            with open(doc_xml_path, 'r', encoding='utf-8') as f:
                document_xml = f.read()

            print("Replacing cover page placeholders...")
            language = variables.get("language", "en")
            document_xml = replace_cover_placeholders(document_xml, variables)

            # Generate body content XML
            print("Generating body content from sections...")
            body_content_xml = generate_body_content_xml(content_data, language)

            # NOTE: Architecture diagram is inserted AFTER the DOCX is built,
            # using python-docx's new_pic_inline() for reliable image embedding.
            # Raw OOXML injection was unreliable (Word rejected the files).
            if arch_diagram_path and not os.path.exists(arch_diagram_path):
                print(f"Warning: Architecture diagram not found: {arch_diagram_path}")
                arch_diagram_path = None  # skip insertion later

            # Replace the body content markers (<!-- BODY_CONTENT_START --> to <!-- BODY_CONTENT_END -->)
            # with the generated content
            body_marker_pattern = re.compile(
                r'<!-- BODY_CONTENT_START -->\s*<!-- BODY_CONTENT_END -->',
                re.DOTALL
            )

            if body_marker_pattern.search(document_xml):
                document_xml = body_marker_pattern.sub(body_content_xml, document_xml)
                print("Injected body content via comment markers")
            else:
                print("Warning: Could not find BODY_CONTENT markers in document.xml")
                return False

            # Write modified document.xml
            with open(doc_xml_path, 'w', encoding='utf-8') as f:
                f.write(document_xml)

            # Repack into DOCX (ZIP) file
            print(f"Repacking into DOCX file: {output_path}")

            # Ensure output directory exists
            os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

            # Create ZIP file with DEFLATED compression
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as docx_zip:
                for root, dirs, files in os.walk(working_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        # Calculate relative path from working directory
                        arcname = os.path.relpath(file_path, working_dir)
                        docx_zip.write(file_path, arcname)

            print(f"Successfully created DOCX: {output_path}")

            # --- Insert architecture diagram using python-docx ---
            if arch_diagram_path and os.path.exists(arch_diagram_path):
                print(f"Embedding architecture diagram via python-docx: {arch_diagram_path}")
                _insert_arch_diagram_with_docx(output_path, arch_diagram_path)

            return True

    except Exception as e:
        print(f"Error processing template: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return False


def _insert_arch_diagram_with_docx(docx_path: str, image_path: str) -> None:
    """
    Open the generated DOCX with python-docx and insert the architecture
    diagram image into the 'Architecture Diagram' section.

    This uses python-docx's ``new_pic_inline()`` which correctly manages
    relationships, content types, and OOXML structure — unlike raw XML
    injection which Word may reject.
    """
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = DocxDocument(docx_path)

    # Find the Architecture Diagram heading
    target_para = None
    for i, p in enumerate(doc.paragraphs):
        if 'Architecture' in p.text and 'Diagram' in p.text and 'Heading' in p.style.name:
            # Use the paragraph right after the heading (description text)
            if i + 1 < len(doc.paragraphs):
                target_para = doc.paragraphs[i + 1]
            else:
                target_para = p
            break

    if target_para is None:
        print("Warning: Could not find 'Architecture Diagram' heading — image not inserted")
        return

    # Calculate width: max 6 inches, preserve aspect ratio
    width_inches = 6.0
    if HAS_PIL:
        try:
            with Image.open(image_path) as img:
                w_px, h_px = img.size
                dpi = img.info.get('dpi', (150, 150))
                dpi_x = dpi[0] if dpi[0] > 0 else 150
                w_in = w_px / dpi_x
                if w_in < width_inches:
                    width_inches = w_in
        except Exception:
            pass

    # Create a new centered paragraph after the description
    new_para = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    new_para.append(pPr)
    target_para._element.addnext(new_para)

    # Use python-docx to build the inline image element
    inline = doc.part.new_pic_inline(image_path, width=Inches(width_inches))

    run = OxmlElement('w:r')
    drawing = OxmlElement('w:drawing')
    drawing.append(inline)
    run.append(drawing)
    new_para.append(run)

    doc.save(docx_path)
    print(f"Architecture diagram inserted successfully ({width_inches:.1f}\" wide)")


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description="Generate scope documents from a DOCX template"
    )
    parser.add_argument(
        "--template-dir",
        required=True,
        help="Path to unpacked DOCX template directory"
    )
    parser.add_argument(
        "--variables",
        required=True,
        help="Path to variables.json file"
    )
    parser.add_argument(
        "--content",
        required=True,
        help="Path to content.json file"
    )
    parser.add_argument(
        "--output",
        required=True,
        help="Output DOCX file path"
    )
    parser.add_argument(
        "--arch-diagram",
        help="Optional path to architecture diagram image"
    )

    args = parser.parse_args()

    # Validate input files exist
    if not os.path.isdir(args.template_dir):
        print(f"Error: Template directory not found: {args.template_dir}", file=sys.stderr)
        return 1

    if not os.path.isfile(args.variables):
        print(f"Error: Variables file not found: {args.variables}", file=sys.stderr)
        return 1

    if not os.path.isfile(args.content):
        print(f"Error: Content file not found: {args.content}", file=sys.stderr)
        return 1

    # Load JSON files
    try:
        with open(args.variables, 'r', encoding='utf-8') as f:
            variables = json.load(f)
        print(f"Loaded variables from {args.variables}")
    except json.JSONDecodeError as e:
        print(f"Error parsing variables.json: {e}", file=sys.stderr)
        return 1

    try:
        with open(args.content, 'r', encoding='utf-8') as f:
            content_data = json.load(f)
        print(f"Loaded content from {args.content}")
    except json.JSONDecodeError as e:
        print(f"Error parsing content.json: {e}", file=sys.stderr)
        return 1

    # Validate arch-diagram if provided
    if args.arch_diagram and not os.path.isfile(args.arch_diagram):
        print(f"Error: Architecture diagram not found: {args.arch_diagram}", file=sys.stderr)
        return 1

    # Process template
    success = process_template(
        args.template_dir,
        variables,
        content_data,
        args.output,
        args.arch_diagram
    )

    return 0 if success else 1


if __name__ == "__main__":
    sys.exit(main())
